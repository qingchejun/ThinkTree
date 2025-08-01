"""
文件解析器模块 - 支持多种文件格式解析
"""

import os
import re
from pathlib import Path
from typing import Optional
import docx
import PyPDF2
import pdfplumber
from io import BytesIO

class FileParser:
    """文件解析器类"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.md', '.docx', '.pdf', '.srt']
    
    def parse_file(self, file_path: str, file_type: str) -> Optional[str]:
        """
        统一文件解析入口
        Args:
            file_path: 文件路径
            file_type: 文件类型扩展名
        Returns:
            解析后的文本内容
        """
        if file_type not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_type}")
        
        try:
            if file_type == '.txt':
                return self.parse_txt(file_path)
            elif file_type == '.md':
                return self.parse_md(file_path)
            elif file_type == '.docx':
                return self.parse_docx(file_path)
            elif file_type == '.pdf':
                return self.parse_pdf(file_path)
            elif file_type == '.srt':
                return self.parse_srt(file_path)
        except Exception as e:
            raise Exception(f"文件解析失败: {str(e)}")
        
        return None
    
    def parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def parse_md(self, file_path: str) -> str:
        """解析Markdown文件，保留结构信息"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
            
            # 简单的Markdown结构提取，保留层次信息
            return self._process_markdown_content(content)
            
        except Exception as e:
            raise Exception(f"Markdown文件解析失败: {str(e)}")
    
    def _process_markdown_content(self, content: str) -> str:
        """处理Markdown内容，保留结构"""
        if not content:
            return ""
        
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 处理标题
            if line.startswith('#'):
                # 保留标题层级信息
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                if title:
                    processed_lines.append(f"{'  ' * (level - 1)}[标题{level}] {title}")
                continue
            
            # 处理列表
            if line.startswith(('-', '*', '+')):
                item = line[1:].strip()
                if item:
                    processed_lines.append(f"• {item}")
                continue
            
            # 处理数字列表
            if re.match(r'^\d+\.', line):
                item = re.sub(r'^\d+\.\s*', '', line)
                if item:
                    processed_lines.append(f"• {item}")
                continue
            
            # 处理代码块（跳过）
            if line.startswith('```'):
                continue
            
            # 处理引用
            if line.startswith('>'):
                quote = line[1:].strip()
                if quote:
                    processed_lines.append(f"「{quote}」")
                continue
            
            # 处理链接和图片（提取文本部分）
            line = re.sub(r'!\[([^\]]*)\]\([^\)]*\)', r'\1', line)  # 图片
            line = re.sub(r'\[([^\]]*)\]\([^\)]*\)', r'\1', line)   # 链接
            
            # 处理加粗、斜体
            line = re.sub(r'\*\*([^\*]+)\*\*', r'\1', line)  # 加粗
            line = re.sub(r'\*([^\*]+)\*', r'\1', line)      # 斜体
            line = re.sub(r'__([^_]+)__', r'\1', line)       # 加粗
            line = re.sub(r'_([^_]+)_', r'\1', line)         # 斜体
            
            # 处理行内代码
            line = re.sub(r'`([^`]+)`', r'\1', line)
            
            # 添加普通文本
            if line.strip():
                processed_lines.append(line.strip())
        
        return '\n'.join(processed_lines)
    
    def parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        doc = docx.Document(file_path)
        text_content = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    def parse_pdf(self, file_path: str) -> str:
        """解析PDF文件，优先使用 pdfplumber，fallback 到 PyPDF2"""
        text_content = []
        
        try:
            # 尝试使用 pdfplumber（更好的文本提取）
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        # 清理文本格式
                        cleaned_text = self._clean_pdf_text(text)
                        if cleaned_text:
                            text_content.append(cleaned_text)
        except Exception:
            # fallback 到 PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            cleaned_text = self._clean_pdf_text(text)
                            if cleaned_text:
                                text_content.append(cleaned_text)
            except Exception as e:
                raise Exception(f"PDF解析失败: {str(e)}")
        
        if not text_content:
            raise Exception("PDF文件没有可提取的文本内容")
        
        return "\n".join(text_content)
    
    def _clean_pdf_text(self, text: str) -> str:
        """清理PDF提取的文本"""
        if not text:
            return ""
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除页眉页脚常见模式（可根据需要调整）
        text = re.sub(r'第\s*\d+\s*页', '', text)
        text = re.sub(r'Page\s*\d+', '', text)
        
        # 修复常见的PDF提取问题
        text = text.replace('ﬁ', 'fi').replace('ﬂ', 'fl')
        
        return text.strip()
    
    def parse_srt(self, file_path: str) -> str:
        """解析SRT字幕文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
        
            # 使用统一的SRT内容处理方法
            return self._process_srt_content(content)
        
        except Exception as e:
            raise Exception(f"SRT文件解析失败: {str(e)}")
    
    def _merge_subtitle_sentences(self, subtitle_texts: list) -> str:
        """合并字幕句子，提高可读性"""
        if not subtitle_texts:
            return ""
        
        merged_sentences = []
        current_sentence = ""
        
        for text in subtitle_texts:
            # 如果当前文本以句号、问号、感叹号结尾，或者是独立的句子
            if current_sentence and (
                text[0].isupper() or  # 新句子（首字母大写）
                current_sentence.endswith(('.', '!', '?', '。', '！', '？'))  # 当前句子已结束
            ):
                if current_sentence.strip():
                    merged_sentences.append(current_sentence.strip())
                current_sentence = text
            else:
                # 继续拼接当前句子
                if current_sentence:
                    current_sentence += " " + text
                else:
                    current_sentence = text
        
        # 添加最后一个句子
        if current_sentence.strip():
            merged_sentences.append(current_sentence.strip())
        
        return "\n".join(merged_sentences)
    
    def parse_from_bytes(self, file_bytes: bytes, filename: str) -> Optional[str]:
        """
        从字节流解析文件内容（优化版本，避免临时文件IO）
        Args:
            file_bytes: 文件字节内容
            filename: 文件名
        Returns:
            解析后的文本内容
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        try:
            if file_ext == '.txt':
                return self._parse_txt_from_bytes(file_bytes)
            elif file_ext == '.md':
                return self._parse_md_from_bytes(file_bytes)
            elif file_ext == '.docx':
                return self._parse_docx_from_bytes(file_bytes, filename)
            elif file_ext == '.pdf':
                return self._parse_pdf_from_bytes(file_bytes, filename)
            elif file_ext == '.srt':
                return self._parse_srt_from_bytes(file_bytes)
        except Exception as e:
            raise Exception(f"文件解析失败: {str(e)}")
        
        return None
    
    def _parse_txt_from_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析纯文本文件"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        # 如果所有编码都失败，使用utf-8并忽略错误
        return file_bytes.decode('utf-8', errors='ignore')
    
    def _parse_md_from_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析Markdown文件"""
        # 先解码为文本
        content = self._parse_txt_from_bytes(file_bytes)
        # 处理Markdown结构
        return self._process_markdown_content(content)
    
    def _parse_docx_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """从字节流解析DOCX文件"""
        try:
            from io import BytesIO
            from docx import Document
            
            # 使用BytesIO避免临时文件
            doc_stream = BytesIO(file_bytes)
            doc = Document(doc_stream)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 解析表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            # 如果内存解析失败，降级到临时文件方式
            return self._fallback_to_temp_file(file_bytes, filename, '.docx')
    
    def _parse_pdf_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """从字节流解析PDF文件"""
        try:
            from io import BytesIO
            import pdfplumber
            
            # 使用BytesIO避免临时文件
            pdf_stream = BytesIO(file_bytes)
            
            with pdfplumber.open(pdf_stream) as pdf:
                text_content = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        # 清理文本
                        cleaned_text = self._clean_pdf_text(page_text)
                        if cleaned_text.strip():
                            text_content.append(cleaned_text)
                
                return "\n".join(text_content)
                
        except Exception as e:
            try:
                # 降级到PyPDF2
                from io import BytesIO
                import PyPDF2
                
                pdf_stream = BytesIO(file_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_stream)
                
                text_content = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        cleaned_text = self._clean_pdf_text(page_text)
                        if cleaned_text.strip():
                            text_content.append(cleaned_text)
                
                return "\n".join(text_content)
                
            except Exception as e2:
                # 如果内存解析都失败，降级到临时文件方式
                return self._fallback_to_temp_file(file_bytes, filename, '.pdf')
    
    def _process_srt_content(self, content: str) -> str:
        """处理SRT格式内容，提取字幕文本"""
        try:
            # 解析SRT格式，提取字幕文本
            subtitle_texts = []
            lines = content.split('\n')
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # 跳过空行
                if not line:
                    i += 1
                    continue
                
                # 跳过序号行（纯数字）
                if line.isdigit():
                    i += 1
                    continue
                
                # 跳过时间戳行（包含 -->）
                if '-->' in line:
                    i += 1
                    continue
                
                # 这是字幕文本行
                if line:
                    # 清理HTML标签（如果有）
                    cleaned_line = re.sub(r'<[^>]+>', '', line)
                    if cleaned_line.strip():
                        subtitle_texts.append(cleaned_line.strip())
                
                i += 1
            
            if not subtitle_texts:
                raise Exception("SRT文件中没有找到有效的字幕文本")
            
            # 合并重复的短句，提高可读性
            return self._merge_subtitle_sentences(subtitle_texts)
        except Exception as e:
            raise Exception(f"SRT内容处理失败: {str(e)}")
    
    def _parse_srt_from_bytes(self, file_bytes: bytes) -> str:
        """从字节流解析SRT字幕文件"""
        # 先解码为文本
        content = self._parse_txt_from_bytes(file_bytes)
        # 处理SRT格式
        return self._process_srt_content(content)
    
    def _fallback_to_temp_file(self, file_bytes: bytes, filename: str, file_ext: str) -> str:
        """降级到临时文件方式（仅在内存解析失败时使用）"""
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            result = self.parse_file(temp_path, file_ext)
            return result
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)

# 创建全局文件解析器实例
file_parser = FileParser()